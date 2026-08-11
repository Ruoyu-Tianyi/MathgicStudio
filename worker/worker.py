"""Mathgic Studio Worker —— 任务队列消费者。

轮询 Supabase jobs 表，领取 pending 任务，调用 Kimi API 按五阶段工作流
（P0 审题 → P1 数据方案 → P2 建模求解 → P3 论文写作 → P4 校验交付）生成论文，
产出 docx 回传 Supabase Storage 并更新任务状态。

运行：python worker.py   （配置见 .env）
"""
import os
import re
import sys
import time
import uuid
import json
import datetime

import requests
from docx import Document
from docx.shared import Pt

# ---------- 配置 ----------

def load_env(path):
    if not os.path.exists(path):
        return
    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())

load_env(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

SUPABASE_URL = os.environ.get("SUPABASE_URL", "").rstrip("/")
SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
MOONSHOT_KEY = os.environ.get("MOONSHOT_API_KEY", "")
MODEL = os.environ.get("MODEL", "kimi-k2-0905-preview")
POLL_INTERVAL = int(os.environ.get("POLL_INTERVAL", "10"))

if not (SUPABASE_URL and SERVICE_KEY and MOONSHOT_KEY):
    sys.exit("缺少配置：请复制 .env.example 为 .env 并填入 SUPABASE_URL / "
             "SUPABASE_SERVICE_ROLE_KEY / MOONSHOT_API_KEY")

SB_HEADERS = {
    "apikey": SERVICE_KEY,
    "Authorization": f"Bearer {SERVICE_KEY}",
    "Content-Type": "application/json",
    "Prefer": "return=representation",
}
KIMI_URL = "https://api.moonshot.cn/v1/chat/completions"

CONTEST_NAME = {"cumcm": "全国大学生数学建模竞赛（CUMCM）", "mcm": "美国大学生数学建模竞赛（MCM/ICM）"}
CONTEST_LANG = {"cumcm": "中文（术语保留英文）", "mcm": "英文"}

# ---------- Supabase 操作 ----------

def sb_get(path):
    r = requests.get(f"{SUPABASE_URL}/rest/v1/{path}", headers=SB_HEADERS, timeout=30)
    r.raise_for_status()
    return r.json()


def sb_patch(table, job_id, payload):
    r = requests.patch(f"{SUPABASE_URL}/rest/v1/{table}?id=eq.{job_id}",
                       headers=SB_HEADERS, json=payload, timeout=30)
    r.raise_for_status()
    return r.json()


def sb_upload(path, data, content_type):
    r = requests.post(f"{SUPABASE_URL}/storage/v1/object/jobs/{path}",
                      headers={"apikey": SERVICE_KEY,
                               "Authorization": f"Bearer {SERVICE_KEY}",
                               "Content-Type": content_type,
                               "x-upsert": "true"},
                      data=data, timeout=120)
    r.raise_for_status()


def sb_download(path):
    r = requests.get(f"{SUPABASE_URL}/storage/v1/object/jobs/{path}",
                     headers=SB_HEADERS, timeout=60)
    r.raise_for_status()
    return r.content

# ---------- Kimi API ----------

def chat(system, user, max_retries=3):
    body = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.6,
    }
    for attempt in range(max_retries):
        try:
            r = requests.post(
                KIMI_URL,
                headers={"Authorization": f"Bearer {MOONSHOT_KEY}",
                         "Content-Type": "application/json"},
                json=body, timeout=600)
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"  [kimi] 第 {attempt + 1} 次调用失败: {e}", flush=True)
            if attempt == max_retries - 1:
                raise
            time.sleep(5 * (attempt + 1))

# ---------- 生成管线（复刻 skill/ 五阶段） ----------

SYS = ("你是数学建模竞赛特等奖级别的解题与写作专家，严格执行 math-modeling-contest "
       "工作流：先推导后写作、不编造数据、模型必有检验与灵敏度分析、公式用 LaTeX。"
       "输出只给正文内容，不要解释你做了什么。")


def stage_p0(job, problem):
    """P0 审题立项：赛道判定 + 问题拆解 + 技术路线"""
    user = f"""【{CONTEST_NAME[job['contest']]}赛题】
{problem}

完成 P0 审题立项，输出《问题分析》：
1. 赛道判定：逐问统计 C 型（数据驱动）信号 vs B 型（机理/优化）信号，给出判定结论与理由。
2. 把赛题拆解为子问题 Q1…Qn，每问标注：题型（评价/预测/优化/机理/数据分析）、拟用模型候选、所需数据及来源。
3. 整体技术路线说明。"""
    return chat(SYS, user)


def stage_p1_p2(job, problem, p0):
    """P1+P2 数据方案与建模推导（推导稿先行，含六检查与灵敏度设计）"""
    user = f"""【赛题】
{problem}

【P0 问题分析】
{p0}

完成 P1 数据方案与 P2 建模推导稿：
1. 数据计划：每问所需数据的真实来源（公开统计/赛题附件/权威数据库）；取不到真实数据的，设计可复现的仿真/假设数据方案，并写明将如何在论文"模型假设"中声明。严禁编造具体数值结果。
2. 每问完成纸面推导：定义与符号 → 模型建立 → 推导 → 可解形式。
3. 每个核心模型做六检查（量纲/退化/不变量/界/良态性/反例），写出结论。
4. 给出每个模型的求解算法步骤与灵敏度分析设计（扰动哪个参数、扫描范围、观察什么指标）。"""
    return chat(SYS, user)


def stage_p3(job, problem, p0, p12):
    """P3 论文写作：完整论文 Markdown"""
    user = f"""【赛题】
{problem}

【P0 问题分析】
{p0}

【P1/P2 数据方案与推导稿】
{p12}

按 {CONTEST_NAME[job['contest']]} 规范撰写完整论文，语言：{CONTEST_LANG[job['contest']]}。
结构：标题 → 摘要（问题→方法→结果→结论，一段式）→ 关键词 → 一、问题重述 →
二、问题分析 → 三、模型假设 → 四、符号说明 → 五、模型建立与求解（分问，含模型检验与
灵敏度分析）→ 六、模型评价（优缺点与改进方向）→ 参考文献 → 附录（核心代码，Python）。
要求：
- 公式用 LaTeX（$...$ / $$...$$），符号全文一致，图表位置用【图 N】/【表 N】占位并配题注。
- 凡涉及具体数值结果而未真实计算的，标注"（待数值实验验证）"，不得编造确定数值。
- 摘要最后打磨，200–400 字（美赛 1 页内）。
- 用 Markdown 标题层级（#, ##, ###）组织全文。"""
    return chat(SYS, user)


def stage_p4(job, paper_md):
    """P4 校验：结构完整性检查"""
    required = ["摘要", "关键词", "问题重述", "问题分析", "模型假设",
                "符号", "模型建立", "模型检验", "灵敏度", "模型评价", "参考文献", "附录"]
    missing = [s for s in required if s not in paper_md]
    return missing

# ---------- Markdown → docx ----------

def md_to_docx(md, title, out_path):
    doc = Document()
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"].font.size = Pt(11)
    for line in md.splitlines():
        s = line.rstrip()
        if not s.strip():
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", s)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            continue
        if re.match(r"^\s*[-*]\s+", s):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", s), style="List Bullet")
            continue
        if re.match(r"^\s*\d+[.、]\s*", s):
            doc.add_paragraph(re.sub(r"^\s*\d+[.、]\s*", "", s), style="List Number")
            continue
        p = doc.add_paragraph(s)
        p.paragraph_format.first_line_indent = Pt(22)
    doc.save(out_path)

# ---------- 主流程 ----------

def read_problem(job):
    parts = []
    if job.get("problem_text"):
        parts.append(job["problem_text"])
    if job.get("problem_file_path"):
        try:
            data = sb_download(job["problem_file_path"])
            parts.append(f"[赛题附件 PDF 已上传（{len(data)} 字节），"
                         "文本提取不可用时请基于题目文本部分作答]")
        except Exception as e:
            parts.append(f"[附件下载失败: {e}]")
    return "\n\n".join(parts)


def process(job):
    jid = job["id"]
    print(f"[{datetime.datetime.now():%H:%M:%S}] 领取任务 {jid[:8]}…", flush=True)
    sb_patch("jobs", jid, {"status": "running", "stage": "P0 审题立项"})

    problem = read_problem(job)
    if not problem.strip():
        raise ValueError("任务缺少赛题内容（problem_text / problem_file_path 均为空）")

    p0 = stage_p0(job, problem)
    sb_patch("jobs", jid, {"stage": "P1/P2 数据方案与建模推导"})
    p12 = stage_p1_p2(job, problem, p0)
    sb_patch("jobs", jid, {"stage": "P3 论文写作"})
    paper = stage_p3(job, problem, p0, p12)

    sb_patch("jobs", jid, {"stage": "P4 校验交付"})
    missing = stage_p4(job, paper)
    if missing:
        print(f"  [P4] 结构缺项: {missing}，补充生成中…", flush=True)
        paper += "\n\n" + chat(SYS, f"论文缺少以下章节：{missing}。请补写这些章节（Markdown 格式）。"
                                     f"\n\n【已有论文】\n{paper[-6000:]}")

    out = f"paper_{jid}.docx"
    md_to_docx(paper, "paper", out)
    with open(out, "rb") as f:
        data = f.read()
    result_path = f"output/{jid}.docx"
    sb_upload(result_path, data,
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    sb_patch("jobs", jid, {"status": "done", "stage": "已完成",
                           "result_file_path": result_path, "error": None})
    os.remove(out)
    print(f"[{datetime.datetime.now():%H:%M:%S}] 任务 {jid[:8]}… 完成 → {result_path}", flush=True)


def main():
    print(f"Mathgic Worker 启动 | 模型 {MODEL} | 每 {POLL_INTERVAL}s 轮询", flush=True)
    while True:
        try:
            jobs = sb_get("jobs?status=eq.pending&order=created_at.asc&limit=1")
            if jobs:
                try:
                    process(jobs[0])
                except Exception as e:
                    print(f"  [失败] {e}", flush=True)
                    sb_patch("jobs", jobs[0]["id"],
                             {"status": "failed", "error": str(e)[:500]})
            else:
                time.sleep(POLL_INTERVAL)
        except KeyboardInterrupt:
            raise
        except Exception as e:
            print(f"  [轮询异常] {e}，{POLL_INTERVAL * 2}s 后重试", flush=True)
            time.sleep(POLL_INTERVAL * 2)


if __name__ == "__main__":
    main()
