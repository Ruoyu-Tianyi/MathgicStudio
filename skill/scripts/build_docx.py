#!/usr/bin/env python3
"""Build a CUMCM-standard .docx paper from paper.md, ON TOP OF the official
template assets/cumcm-template.docx (V1.1).

The template contributes: A4 page & margins (T/B 2.54, L/R 2.70 cm), its own
paragraph styles (Heading 1/2/3 黑体, Normal TNR+宋体 12pt, 图表标题 10.5pt
bold), the 三线表 table style, and the centered page-number footer. This
script empties the template body and re-writes the paper with those styles.

Fully offline: math is rendered with matplotlib mathtext to embedded PNGs
(\tag{n} stripped, re-attached as a right-aligned equation number).

Usage:
    python build_docx.py <paper.md> [--out paper.docx] [--template PATH]

Markdown subset: #..#### headings, paragraphs, **bold**, - and 1. lists,
| tables |, ![caption](path), $$..$$ display math, $..$ inline math.
"""
import argparse
import re
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# OMML (native Word equation) pipeline: LaTeX -> MathML (latex2mathml)
# -> OMML (Office-bundled MML2OMML.XSL). Falls back to mathtext PNG.
import latex2mathml.converter
from lxml import etree

MML2OMML = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_omml_tf = None
if Path(MML2OMML).is_file():
    _omml_tf = etree.XSLT(etree.parse(MML2OMML))
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def preprocess_latex(latex):
    tag = None
    m = TAG_RE.search(latex)
    if m:
        tag = m.group(1)
        latex = TAG_RE.sub("", latex)
    latex = latex.replace(r"\dfrac", r"\frac").strip()
    latex = re.sub(r"\\le(?![a-zA-Z])", r"\\leq", latex)
    latex = re.sub(r"\\ge(?![a-zA-Z])", r"\\geq", latex)
    return latex, tag


def latex_to_omml(latex):
    """Returns m:oMath lxml element, or None on failure."""
    if _omml_tf is None:
        return None
    try:
        mml = latex2mathml.converter.convert(latex)
        omml = _omml_tf(etree.fromstring(mml.encode()))
        return omml.getroot()
    except Exception:
        return None

TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "cumcm-template.docx"

# manual-run fonts for elements the template formats manually (title/摘要头)
TITLE_FONT = {"ea": "黑体", "ascii": "Times New Roman", "size": 16}
ABSH_FONT = {"ea": "黑体", "ascii": "Times New Roman", "size": 14}
BODY_SIZE = 12

MATH_IMG = Path(__file__).resolve().parent / "_math_tmp"
MATH_IMG.mkdir(exist_ok=True)
_math_counter = 0
_eq_counter = 0  # V3.9：display math 全文顺序编号（含附录）


def set_font(run, font, bold=None):
    run.font.name = font["ascii"]
    run.font.size = Pt(font["size"])
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font["ea"])


TAG_RE = re.compile(r"\\tag\{([^}]*)\}")


def math_png(latex, fontsize=12):
    global _math_counter
    latex, tag = preprocess_latex(latex)
    _math_counter += 1
    path = MATH_IMG / f"m{_math_counter}.png"
    fig = plt.figure(figsize=(0.1, 0.1))
    t = fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    fig.canvas.draw()
    bbox = t.get_window_extent()
    plt.close(fig)
    fig = plt.figure(figsize=(max(bbox.width, 2) / 72, max(bbox.height, 2) / 72))
    fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    fig.savefig(path, dpi=300, bbox_inches="tight", pad_inches=0.01, transparent=True)
    plt.close(fig)
    from PIL import Image
    with Image.open(path) as im:
        w, h = im.size
    return str(path), w, h, tag


INLINE_MATH_RE = re.compile(r"\$([^$]+)\$")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
IMG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)(?:\{w=([\d.]+)cm\})?\s*$")


def auto_img_width_cm(img_path):
    """Aspect-ratio tiers: wide >=1.8 -> 13.5 cm, square <=1.2 -> 10 cm,
    else 12 cm. Used when no {w=..cm} annotation is present."""
    from PIL import Image
    with Image.open(img_path) as im:
        w, h = im.size
    ratio = w / max(h, 1)
    if ratio >= 1.8:
        return 13.5
    if ratio <= 1.2:
        return 10.0
    return 12.0
CAPTION_RE = re.compile(r"^(图|表)\s*\d+\s*[:：]")


def add_runs(p, text, base_size=BODY_SIZE):
    """Add runs handling $math$ and **bold**; style/fonts come from paragraph style."""
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            _add_bold_aware(p, text[pos:m.start()])
        latex, _ = preprocess_latex(m.group(1))
        omml = latex_to_omml(latex)
        if omml is not None:  # native Word equation
            p._p.append(omml)
        else:                 # fallback: mathtext PNG
            img, w, h, _ = math_png(m.group(1), fontsize=int(base_size))
            run = p.add_run()
            height_pt = base_size * 1.15
            run.add_picture(img, height=Pt(height_pt), width=Pt(height_pt * w / h))
        pos = m.end()
    if pos < len(text):
        _add_bold_aware(p, text[pos:])


CODE_RE = re.compile(r"`([^`]+)`")


def _add_bold_aware(p, text):
    def emit(seg):
        pos2 = 0
        for cm in CODE_RE.finditer(seg):
            if cm.start() > pos2:
                p.add_run(seg[pos2:cm.start()])
            r = p.add_run(cm.group(1))
            r.font.name = "Consolas"
            pos2 = cm.end()
        if pos2 < len(seg):
            p.add_run(seg[pos2:])
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            emit(text[pos:m.start()])
        r = p.add_run(re.sub(CODE_RE, r"\1", m.group(1)))
        r.bold = True
        pos = m.end()
    if pos < len(text):
        emit(text[pos:])


def set_cell_border(cell, **edges):
    """edges: edge=(on, size_eighth_pt). e.g. top=(True, 12) -> 1.5pt line."""
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        spec = edges.get(edge)
        if spec and spec[0]:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(spec[1]))
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def set_table_cell_margins(t, side_cm=0.2):
    tbl_pr = t._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 40), ("left", int(side_cm * 567)),
                      ("bottom", 40), ("right", int(side_cm * 567))):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def add_table(doc, rows):
    """3-line table per 优秀作品 spec: 1.5pt top/bottom rules, 0.5pt header
    rule; first column narrow+centered, others left-aligned; cells
    vertically centered with 0.2 cm side margins; borders drawn directly
    (table-style tblLook is unreliable).

    Pagination: every row gets cantSplit (a row never breaks across pages);
    tables with ≤30 rows additionally get keepNext on all rows but the last,
    so the whole table moves to the next page as one block instead of being
    split (user rule: 宁可另起一页). Very long tables (>30 rows, e.g. appendix
    data) are allowed to flow across pages, but still only between rows."""
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    KEEP_TOGETHER_MAX_ROWS = 30
    n_rows, n_cols = len(rows), len(rows[0])
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    text_w_cm = 21.0 - 2 * 2.7          # A4 minus template L/R margins
    first_w = 3.0 if n_cols > 1 else text_w_cm
    other_w = (text_w_cm - first_w) / max(n_cols - 1, 1)
    widths = [first_w] + [other_w] * (n_cols - 1)
    set_table_cell_margins(t)
    for i, row in enumerate(rows):
        # 行不跨页
        tr_pr = t.rows[i]._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        tr_pr.append(cant)
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.width = Cm(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, cell_text, base_size=10.5)
            if n_rows <= KEEP_TOGETHER_MAX_ROWS and i < n_rows - 1:
                # 整表不跨页：除末行外所有段落 keepNext，与下一行绑定
                for cp in cell.paragraphs:
                    cp.paragraph_format.keep_with_next = True
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(10.5)
                if i == 0:
                    r.bold = True
            set_cell_border(cell,
                            top=(i == 0, 12),            # 顶线 1.5pt
                            bottom=(True, 4) if i == 0 else
                                   ((True, 12) if i == n_rows - 1 else None),
                            # 栏目线 0.5pt / 底线 1.5pt，其余无线
                            left=None, right=None, insideH=None, insideV=None)
    return t


def first_line_indent(p, chars=2):
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    p._p.get_or_add_pPr().append(ind)


def clear_body(doc):
    body = doc.element.body
    for el in list(body):
        if el.tag != qn("w:sectPr"):
            body.remove(el)


def fix_h1_numbering(doc):
    """Template Heading-1 auto-number renders '一、'; change to '一 ' (space)."""
    for part in doc.part.package.parts:
        if str(part.partname).endswith("numbering.xml"):
            for el in part.element.iter():
                if el.tag == qn("w:lvlText") and el.get(qn("w:val"), "").endswith("、"):
                    el.set(qn("w:val"), el.get(qn("w:val"))[:-1] + " ")


def add_page_field(paragraph):
    """Append a PAGE field to a paragraph."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    fld.append(r)
    paragraph._p.append(fld)


def setup_header_footer(doc, title):
    """Sample-paper style: header = paper title left + page number right +
    bottom rule; template's centered footer page number removed."""
    sec = doc.sections[0]
    header, footer = sec.header, sec.footer
    header.is_linked_to_previous = False
    for p in list(header.paragraphs[1:]):
        p._p.getparent().remove(p._p)
    hp = header.paragraphs[0]
    for r in list(hp.runs):
        r._element.getparent().remove(r._element)
    hp.text = ""
    from docx.enum.text import WD_TAB_ALIGNMENT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.6), WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run(title + "\t")
    add_page_field(hp)
    ppr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "000000"); bottom.set(qn("w:space"), "1")
    pbdr.append(bottom)
    ppr.append(pbdr)
    # clear footer (page number moves to header)
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for child in list(p._p):
            if child.tag != qn("w:pPr"):
                p._p.remove(child)


CODE_KW = ("import", "from", "def", "return", "for", "while", "if", "else",
           "elif", "in", "not", "and", "or", "print", "class", "with", "as",
           "lambda", "None", "True", "False", "try", "except", "raise", "pass")


def add_code_block(doc, code_lines):
    """Fenced code block: line numbers + Consolas 8.5pt + light shading."""
    for n, cl in enumerate(code_lines, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        ppr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F5F5F5")
        ppr.append(shd)
        rno = p.add_run(f"{n:>4}  ")
        rno.font.name = "Consolas"; rno.font.size = Pt(8.5)
        rno.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        if cl.strip().startswith("#"):
            r = p.add_run(cl)
            r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        else:
            r = p.add_run(cl)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)


def append_code_files(doc, code_dir):
    """--appendix-code: append every *.py as Heading 2 + numbered code block."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    files = sorted(Path(code_dir).glob("*.py"))
    if not files:
        print(f"WARN: no .py files in {code_dir}")
        return
    for f in files:
        try:
            p = doc.add_paragraph(style="Heading 2")
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = _AL.LEFT
        p.add_run(f.name)
        lines = f.read_text(encoding="utf-8").splitlines()
        add_code_block(doc, lines)
    print(f"appendix: {len(files)} code file(s) embedded")


def build(md_path: Path, out_path: Path, template: Path, appendix_code=None):
    global _eq_counter
    _eq_counter = 0
    doc = Document(str(template))
    clear_body(doc)
    fix_h1_numbering(doc)

    base = md_path.parent
    lines = md_path.read_text(encoding="utf-8").splitlines()
    paper_title = ""
    seen_abstract = False
    paged_after_abstract = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level, text = len(m.group(1)), m.group(2)
            if level == 1:
                paper_title = text.strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(p, text, base_size=TITLE_FONT["size"])
                for r in p.runs:
                    if r.text:
                        set_font(r, TITLE_FONT)
            elif level == 2 and text.strip() in ("摘要", "摘  要"):
                seen_abstract = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("摘  要")
                set_font(r, ABSH_FONT)
            else:
                if level == 2 and seen_abstract and not paged_after_abstract:
                    doc.add_page_break()  # 摘要独占一页
                    paged_after_abstract = True
                text = re.sub(r"^[一二三四五六七八九十]+、\s*", "", text)
                text = re.sub(r"^\d+(\.\d+)*[、.\s]\s*", "", text)
                style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[level]
                p = doc.add_paragraph(style=style)
                if level == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 范文 H1 居中
                add_runs(p, text)
            i += 1; continue
        if line.strip().startswith("$$"):
            buf = line.strip()[2:]
            while not buf.endswith("$$"):
                i += 1
                buf += lines[i].strip()
            latex, tag = preprocess_latex(buf[:-2])
            omml = latex_to_omml(latex)
            # V3.9：所有 display math 自动编号（\tag{} 可手动覆盖），
            # 公式单行居中、编号右对齐贴页边（(1) (2) ... 全文连续）
            _eq_counter += 1
            if not tag:
                tag = str(_eq_counter)
            p = doc.add_paragraph()
            # V4.2：公式居中 + 编号右对齐的标准制表位方案——段落左对齐，
            # 行首 tab 推到居中制表位（版心中点 7.8cm），公式中心对准版心
            # 中心；第二个 tab 把编号推到右制表位（版心右缘 15.6cm）。
            # V3.9 的段落 CENTER 会把"公式+tab+编号"作为整体居中，公式偏左。
            from docx.enum.text import WD_TAB_ALIGNMENT
            p.paragraph_format.tab_stops.add_tab_stop(Cm(7.8), WD_TAB_ALIGNMENT.CENTER)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(15.6), WD_TAB_ALIGNMENT.RIGHT)
            if omml is not None:
                p.add_run("\t")
                p._p.append(omml)  # inline oMath，跟随制表位布局
                p.add_run("\t（" + tag + "）")
                i += 1; continue
            img, w, h, _ = math_png(buf[:-2], fontsize=13)  # tag 已自取，不覆盖
            p.add_run("\t")
            run = p.add_run()
            max_w_cm, nat_w_cm = 14.0, w / 300 * 2.54
            scale = min(1.0, max_w_cm / max(nat_w_cm, 0.1))
            run.add_picture(img, width=Cm(nat_w_cm * scale), height=Cm(h / 300 * 2.54 * scale))
            p.add_run("\t（" + tag + "）")
            i += 1; continue
        if line.strip().startswith("```"):  # fenced code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # closing fence
            add_code_block(doc, code_lines)
            continue
        m = IMG_RE.match(line.strip())
        if m:
            img_path = (base / m.group(2)).resolve()
            width_cm = float(m.group(3)) if m.group(3) else auto_img_width_cm(img_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img_path), width=Cm(width_cm))
            # 图与"图 N"题注绑定：向后（跳过空行，最多 3 行）找到图题时
            # keep_with_next，图 + 题注同页或整体另起一页（V3.7）
            for lj in range(i + 1, min(i + 4, len(lines))):
                if lines[lj].strip():
                    if (CAPTION_RE.match(lines[lj].strip())
                            and lines[lj].strip().startswith("图")):
                        p.paragraph_format.keep_with_next = True
                    break
            i += 1; continue
        if line.strip().startswith("|"):
            rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c or "---") for c in cells):
                    rows.append(cells)
                j += 1
            # V3.9：真表格至少有表头+分隔行两行；单行 |..| 按正文处理
            # （如 |ρ| 绝对值写法，避免被吞成一行两列的畸形表）
            if rows and j - i >= 2:
                add_table(doc, rows)
                i = j
                continue
        if CAPTION_RE.match(line.strip()):
            cap = re.sub(r"^(图|表)(\s*\d+)\s*[：:]\s*", r"\1\2: ", line.strip())  # 半角冒号
            try:
                p = doc.add_paragraph(style="图表标题")
            except KeyError:
                p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, cap, base_size=10.5)
            if cap.startswith("表"):
                # V3.9 题注间距：表题在表上方，段前 6pt 与正文隔开、段后 3pt 贴表
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                # 表题与表体绑定：题注随表一起另起一页，不留在上一页页尾
                p.paragraph_format.keep_with_next = True
            else:
                # V3.9 题注间距：图题在图下方，段前 3pt 贴图、段后 6pt 与正文隔开
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
            i += 1; continue
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            try:
                p = doc.add_paragraph(style="List Paragraph")
            except KeyError:
                p = doc.add_paragraph()
            add_runs(p, m.group(2) + " " + m.group(3))
            i += 1; continue
        text_line = line.strip()
        if text_line.startswith("**关键词**"):  # 关键词内容空格分隔
            text_line = re.sub(r"[；;,，]\s*", " ", text_line)
        p = doc.add_paragraph()
        first_line_indent(p, 2)
        add_runs(p, text_line)
        i += 1

    if appendix_code:
        append_code_files(doc, appendix_code)
    if paper_title:
        setup_header_footer(doc, paper_title)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("paper", help="path to paper.md")
    ap.add_argument("--out", default=None)
    ap.add_argument("--template", default=str(TEMPLATE))
    ap.add_argument("--appendix-code", default=None, metavar="DIR",
                    help="append all .py files under DIR as appendix code blocks")
    args = ap.parse_args()
    md = Path(args.paper)
    out = Path(args.out) if args.out else md.with_suffix(".docx")
    build(md, out, Path(args.template), appendix_code=args.appendix_code)
    print(f"OK: {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    sys.exit(main())
